from docx import Document
from docx.shared import Inches

def create_tutorial():
    document = Document()
    document.add_heading('RefLookUp: Managing Citations for Biosketches', 0)

    document.add_paragraph(
        'Use RefLookUp to verify your publications and easily import them into MyNCBI for your NIH Biosketch.'
    )

    # Step 1
    document.add_heading('1. Export References', level=1)
    document.add_paragraph(
        'Export your selected references from Google Scholar (or another manager) as a .bib or .csv file.'
    )
    document.add_picture('Tutorial/googlescholar.png', width=Inches(5.5))

    # Step 2
    document.add_heading('2. Verify in RefLookUp', level=1)
    p = document.add_paragraph('Go to the ')
    p.add_run('RefLookUp Web App').bold = True
    p.add_run(' (https://reflookup.streamlit.app/), drag and drop your exported file into the upload field, and click ')
    p.add_run('Search References').bold = True
    p.add_run('.')
    document.add_paragraph(
        'Once finished, look for the RIS download button. This format is specifically designed for citation managers.'
    )
    document.add_picture('Tutorial/app.png', width=Inches(5.5))

    # Step 3
    document.add_heading('3. Import to MyNCBI', level=1)
    document.add_paragraph(
        'Go to your MyNCBI Bibliography. Click "Manage My Bibliography" -> "Add Citations" -> "From a file". Upload the .ris file you just downloaded.'
    )
    document.add_picture('Tutorial/myNCBI.png', width=Inches(5.5))

    document.save('Tutorial/RefLookUp_Tutorial.docx')
    print("Tutorial generated successfully at Tutorial/RefLookUp_Tutorial.docx")

if __name__ == "__main__":
    create_tutorial()
